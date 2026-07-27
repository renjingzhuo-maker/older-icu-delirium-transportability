from __future__ import annotations

import csv
import json
import math
import re
import sys
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
HERE = Path(__file__).resolve().parent
ASSETS = HERE / "assets"
MODELS = ROOT / "results" / "models"
GBTM = ROOT / "results" / "gbtm"
QA = ROOT / "results" / "qa"
TEMPLATE = HERE / "MANUSCRIPT_TEMPLATE.md"
MARKDOWN_OUT = HERE / "MANUSCRIPT_DRAFT.md"
DOCX_OUT = HERE / "MANUSCRIPT_DRAFT.docx"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_DXA = 80
CELL_MARGIN_BOTTOM_DXA = 80
CELL_MARGIN_START_DXA = 120
CELL_MARGIN_END_DXA = 120
HEADER_FILL = "F4F6F9"
BLUE = "2E74B5"
DARK = "20242A"


def read_rows(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def read_one(path: Path) -> dict[str, str]:
    rows = read_rows(path)
    if len(rows) != 1:
        raise ValueError(f"Expected one row in {path}, found {len(rows)}")
    return rows[0]


def f(value: str | float, digits: int = 3) -> str:
    return f"{float(value):.{digits}f}"


def pct(value: str | float, digits: int = 2) -> str:
    return f"{100 * float(value):.{digits}f}"


def metric_ci(row: dict[str, str], metric: str, digits: int = 3) -> str:
    return (
        f"{float(row[metric]):.{digits}f} "
        f"({float(row[f'{metric}_ci_low']):.{digits}f}-"
        f"{float(row[f'{metric}_ci_high']):.{digits}f})"
    )


def find_performance(rows: list[dict[str, str]], prefix: str, model: str) -> dict[str, str]:
    matches = [
        row
        for row in rows
        if row["dataset"].startswith(prefix) and row["dataset"].endswith(model)
    ]
    if len(matches) != 1:
        raise ValueError(f"Expected one performance row for {prefix=} {model=}; found {len(matches)}")
    return matches[0]


def top_feature_label(name: str) -> str:
    labels = {
        "mechvent_24h": "mechanical ventilation",
        "gcs_min": "minimum GCS",
        "rass_min": "minimum RASS",
        "rass_max": "maximum RASS",
        "urineoutput_24h": "24-hour urine output",
        "psychiatric_disorder": "psychiatric disorder",
        "glucose_lab_max": "maximum glucose",
        "bun_max": "maximum blood urea nitrogen",
        "creatinine_max": "maximum creatinine",
        "heart_rate_max": "maximum heart rate",
        "age": "age",
        "cerebrovascular_disease": "cerebrovascular disease",
        "spo2_min": "minimum oxygen saturation",
        "icu_type_Cardiac": "cardiac ICU type",
        "icu_type_Medical": "medical ICU type",
        "icu_type_Surgical": "surgical ICU type",
        "icu_type_Neurologic": "neurologic ICU type",
    }
    return labels.get(name, name.replace("_", " "))


def get_baseline_value(rows: list[dict[str, str]], characteristic: str, database: str) -> str:
    for row in rows:
        if row["Characteristic"] == characteristic:
            return row[database]
    raise KeyError(characteristic)


def build_context() -> dict[str, str]:
    counts = json.loads((MODELS / "analysis_counts.json").read_text(encoding="utf-8"))
    performance = read_rows(MODELS / "all_harmonized_model_performance.csv")
    primary_performance_path = (
        MODELS / "primary_harmonized_model_performance.csv"
    )
    primary_performance = (
        read_rows(primary_performance_path)
        if primary_performance_path.exists() else performance
    )
    clinical_internal = find_performance(primary_performance, "MIMIC internal OOF", "clinical_baseline_harmonized")
    clinical_external = find_performance(primary_performance, "eICU external", "clinical_baseline_harmonized")
    enhanced_internal = find_performance(primary_performance, "MIMIC internal OOF", "nursing_enhanced_harmonized")
    enhanced_external = find_performance(primary_performance, "eICU external", "nursing_enhanced_harmonized")
    no_antipsychotic = find_performance(
        performance,
        "eICU external",
        "nursing_without_antipsychotic_harmonized",
    )
    coding_sensitivity = read_rows(
        MODELS / "coding_harmonization_sensitivity_performance.csv"
    )
    no_psychiatric_external = find_performance(
        coding_sensitivity,
        "eICU external",
        "nursing_without_psychiatric_disorder_harmonized",
    )
    no_icu_type_external = find_performance(
        coding_sensitivity,
        "eICU external",
        "nursing_without_icu_type_harmonized",
    )
    no_race_external = find_performance(
        coding_sensitivity,
        "eICU external",
        "nursing_without_race_harmonized",
    )
    coding_sensitivity_differences = {
        row["omitted_feature"]: row
        for row in read_rows(
            MODELS / "coding_harmonization_sensitivity_paired_bootstrap.csv"
        )
    }
    missing_indicator = find_performance(
        performance,
        "eICU external",
        "nursing_with_missing_indicators",
    )
    full_indicator = find_performance(
        performance,
        "eICU external",
        "nursing_full_with_missing_indicators_harmonized",
    )
    bedside = find_performance(performance, "eICU external", "bedside_score_harmonized")

    delta_rows = read_rows(MODELS / "nursing_increment_paired_bootstrap.csv")
    delta_internal = next(row for row in delta_rows if row["dataset"].startswith("MIMIC"))
    delta_external = next(row for row in delta_rows if row["dataset"].startswith("eICU"))

    flow_rows = read_rows(ASSETS / "cohort_flow_counts.csv")
    mimic_flow = next(row for row in flow_rows if row["source"] == "MIMIC-IV")
    eicu_flow = next(row for row in flow_rows if row["source"] == "eICU-CRD")
    baseline_rows = read_rows(ASSETS / "table_1_baseline_characteristics.csv")

    ablation = read_rows(MODELS / "assessment_selection_model_ablation.csv")
    patient_only = next(
        row for row in ablation
        if row["dataset"] == "eicu" and row["model_variant"] == "patient_features_only"
    )
    hospital_only = next(row for row in ablation if row["model_variant"] == "hospital_only")
    combined = next(row for row in ablation if row["model_variant"] == "patient_features_plus_hospital")
    hospital_attributes = next(
        row for row in ablation
        if row["model_variant"] == "hospital_attributes_only"
    )
    combined_attributes = next(
        row for row in ablation
        if row["model_variant"] == "patient_features_plus_hospital_attributes"
    )
    mimic_selection = next(
        row for row in ablation
        if row["dataset"] == "mimic" and row["model_variant"] == "patient_features_only"
    )

    selection_diagnostics = read_rows(MODELS / "assessment_selection_diagnostics.csv")
    mimic_diagnostics = next(row for row in selection_diagnostics if row["dataset"] == "mimic")
    eicu_diagnostics = next(row for row in selection_diagnostics if row["dataset"] == "eicu")
    ipw_rows = read_rows(MODELS / "assessment_ipw_performance.csv")
    eicu_ipw = next(row for row in ipw_rows if row["dataset"].startswith("eICU"))

    permutation = read_rows(MODELS / "eicu_assessment_selection_permutation_importance.csv")
    hospital_permutation = next(row for row in permutation if row["feature"] == "hospital_site")
    patient_permutation = [
        float(row["mean_auroc_decrease"])
        for row in permutation
        if row["feature"] != "hospital_site"
    ]

    hospital_summary = read_one(MODELS / "eicu_hospital_representation_summary.csv")
    shap_rows = read_rows(MODELS / "nursing_enhanced_harmonized_shap_importance.csv")
    increment_shap = {
        row["feature"]: float(row["mean_abs_shap"])
        for row in read_rows(MODELS / "nursing_assessment_treatment_increment_shap.csv")
    }

    profile = {
        row["feature"]: row
        for row in read_rows(MODELS / "harmonized_feature_profile.csv")
    }
    urine_missing = float(profile["urineoutput_24h"]["eicu_missing"])

    broad = read_one(MODELS / "broad_missing_as_negative_performance.csv")
    recalibration = {
        row["recalibration_method"]: row
        for row in read_rows(
            MODELS
            / "nursing_enhanced_harmonized_recalibration_diagnostics.csv"
        )
    }
    endpoint_performance = read_rows(
        MODELS / "alternative_endpoint_performance.csv"
    )
    endpoint_counts = read_rows(
        MODELS / "alternative_endpoint_counts.csv"
    )
    logistic_performance = read_rows(
        MODELS / "regularized_logistic_benchmark_performance.csv"
    )
    logistic_clinical_external = find_performance(
        logistic_performance,
        "eICU external",
        "clinical_logistic_harmonized",
    )
    logistic_enhanced_external = find_performance(
        logistic_performance,
        "eICU external",
        "nursing_enhanced_logistic_harmonized",
    )

    def endpoint_row(endpoint: str, model: str, prefix: str) -> dict[str, str]:
        matches = [
            row for row in endpoint_performance
            if row["endpoint"] == endpoint
            and row["dataset"].startswith(prefix)
            and model in row["dataset"]
        ]
        if len(matches) != 1:
            raise ValueError(
                f"Expected one endpoint row for {endpoint}, {model}, {prefix}"
            )
        return matches[0]

    strict_enhanced_external = endpoint_row(
        "strict_two_positive_days",
        "nursing_enhanced_strict_two_positive_days_harmonized",
        "eICU external",
    )
    any_enhanced_external = endpoint_row(
        "any_post24_delirium",
        "nursing_enhanced_any_post24_delirium_harmonized",
        "eICU external",
    )
    endpoint_count_map = {
        (row["endpoint"], row["dataset"]): row
        for row in endpoint_counts
    }
    trajectory_performance = read_rows(MODELS / "multiclass_trajectory_performance.csv")
    trajectory_external = next(
        row for row in trajectory_performance if row["dataset"].startswith("eICU")
    )
    trajectory_selection = read_rows(GBTM / "gbtm_model_selection.csv")
    trajectory_two = next(row for row in trajectory_selection if row["G"] == "2")
    trajectory_class = next(
        row
        for row in read_rows(GBTM / "mimic_trajectory_class_summary.csv")
        if row["high_risk_trajectory"] == "1"
    )

    params = json.loads(
        (MODELS / "nursing_enhanced_harmonized_best_parameters.json").read_text(
            encoding="utf-8"
        )
    )
    parameter_text = (
        f"{params['model__n_estimators']} trees, maximum depth "
        f"{params['model__max_depth']}, learning rate "
        f"{params['model__learning_rate']}, minimum child weight "
        f"{params['model__min_child_weight']}, row subsampling "
        f"{params['model__subsample']}, column subsampling "
        f"{params['model__colsample_bytree']}, L1 regularization "
        f"{params['model__reg_alpha']}, and L2 regularization "
        f"{params['model__reg_lambda']}"
    )

    top_features = ", ".join(top_feature_label(row["feature"]) for row in shap_rows[:7])
    mimic_rate = 100 * counts["mimic_late_persistent_events"] / counts["mimic_n"]
    eicu_rate = 100 * counts["eicu_late_persistent_events"] / counts["eicu_n"]
    trajectory_rate = 100 * int(float(trajectory_class["class_n"])) / counts["mimic_n"]

    context = {
        "MIMIC_N": f"{counts['mimic_n']:,}",
        "MIMIC_EVENTS": f"{counts['mimic_late_persistent_events']:,}",
        "MIMIC_RATE": f"{mimic_rate:.2f}",
        "EICU_N": f"{counts['eicu_n']:,}",
        "EICU_EVENTS": f"{counts['eicu_late_persistent_events']:,}",
        "EICU_RATE": f"{eicu_rate:.2f}",
        "MIMIC_BASELINE_INELIGIBLE": f"{int(mimic_flow['baseline_positive_or_other_ineligible']):,}",
        "MIMIC_NO_NEGATIVE": f"{int(mimic_flow['no_documented_negative_baseline']):,}",
        "MIMIC_INSUFFICIENT": f"{int(mimic_flow['fewer_than_2_observed_outcome_days']):,}",
        "EICU_BASELINE_INELIGIBLE": f"{int(eicu_flow['baseline_positive_or_other_ineligible']):,}",
        "EICU_NO_NEGATIVE": f"{int(eicu_flow['no_documented_negative_baseline']):,}",
        "EICU_INSUFFICIENT": f"{int(eicu_flow['fewer_than_2_observed_outcome_days']):,}",
        "MIMIC_AGE": get_baseline_value(baseline_rows, "Age, years", "MIMIC-IV"),
        "EICU_AGE": get_baseline_value(baseline_rows, "Age, years", "eICU-CRD"),
        "EICU_URINE_MISSING": f"{100 * urine_missing:.2f}",
        "CLINICAL_FEATURE_COUNT": str(counts["clinical_feature_count"]),
        "ENHANCED_FEATURE_COUNT": str(counts["nursing_feature_count"]),
        "CLIN_INT_AUROC_CI": metric_ci(clinical_internal, "auroc"),
        "CLIN_EXT_AUROC_CI": metric_ci(clinical_external, "auroc"),
        "CLIN_INT_AUPRC_CI": metric_ci(clinical_internal, "auprc"),
        "CLIN_EXT_AUPRC_CI": metric_ci(clinical_external, "auprc"),
        "CLIN_INT_BRIER_CI": metric_ci(clinical_internal, "brier"),
        "CLIN_EXT_BRIER_CI": metric_ci(clinical_external, "brier"),
        "ENH_INT_AUROC_CI": metric_ci(enhanced_internal, "auroc"),
        "ENH_EXT_AUROC_CI": metric_ci(enhanced_external, "auroc"),
        "ENH_EXT_AUPRC_CI": metric_ci(enhanced_external, "auprc"),
        "DELTA_MIMIC": f"{float(delta_internal['auroc_difference']):.3f}",
        "DELTA_EICU": f"{float(delta_external['auroc_difference']):.3f}",
        "DELTA_MIMIC_FULL": (
            f"{float(delta_internal['auroc_difference']):.3f} "
            f"(95% CI {float(delta_internal['difference_ci_low']):.3f}-"
            f"{float(delta_internal['difference_ci_high']):.3f})"
        ),
        "DELTA_EICU_FULL": (
            f"{float(delta_external['auroc_difference']):.3f} "
            f"(95% CI {float(delta_external['difference_ci_low']):.3f}-"
            f"{float(delta_external['difference_ci_high']):.3f})"
        ),
        "ENH_EXT_CAL_INTERCEPT": f"{float(enhanced_external['calibration_intercept']):.3f}",
        "ENH_EXT_CAL_SLOPE": f"{float(enhanced_external['calibration_slope']):.3f}",
        "ENH_EXT_CAL_INTERCEPT_CI": metric_ci(
            enhanced_external, "calibration_intercept"
        ),
        "ENH_EXT_CAL_SLOPE_CI": metric_ci(
            enhanced_external, "calibration_slope"
        ),
        "ENH_EXT_ECE": f"{float(enhanced_external['ece_10']):.3f}",
        "ENH_EXT_ECE_CI": metric_ci(enhanced_external, "ece_10"),
        "ENH_EXT_AUROC_POINT": f"{float(enhanced_external['auroc']):.3f}",
        "ENH_THRESHOLD": f"{float(enhanced_internal['threshold']):.3f}",
        "ENH_EXT_SENSITIVITY": f"{float(enhanced_external['sensitivity']):.3f}",
        "ENH_EXT_SPECIFICITY": f"{float(enhanced_external['specificity']):.3f}",
        "ENH_EXT_PPV": f"{float(enhanced_external['ppv']):.3f}",
        "ENH_EXT_NPV": f"{float(enhanced_external['npv']):.3f}",
        "FINAL_MODEL_PARAMETERS": parameter_text,
        "TOP_SHAP_FEATURES": top_features,
        "GCS_SHAP": f"{increment_shap.get('gcs_min', math.nan):.3f}",
        "RASS_MIN_SHAP": f"{increment_shap.get('rass_min', math.nan):.3f}",
        "RASS_MAX_SHAP": f"{increment_shap.get('rass_max', math.nan):.3f}",
        "SEDATIVE_SHAP": f"{increment_shap.get('sedative_24h', math.nan):.3f}",
        "MIMIC_SELECTION_RATE": f"{100 * float(mimic_selection['selection_rate']):.2f}",
        "EICU_SELECTION_RATE": f"{100 * float(patient_only['selection_rate']):.2f}",
        "EICU_PATIENT_ONLY_AUC": f"{float(patient_only['auroc']):.3f}",
        "EICU_HOSPITAL_ONLY_AUC": f"{float(hospital_only['auroc']):.3f}",
        "EICU_HOSPITAL_ATTRIBUTES_AUC": f"{float(hospital_attributes['auroc']):.3f}",
        "EICU_PATIENT_ATTRIBUTES_AUC": f"{float(combined_attributes['auroc']):.3f}",
        "EICU_COMBINED_SELECTION_AUC": f"{float(combined['auroc']):.3f}",
        "HOSPITAL_PERMUTATION_DROP": f"{float(hospital_permutation['mean_auroc_decrease']):.3f}",
        "MAX_PATIENT_PERMUTATION_DROP": f"{max(patient_permutation):.3f}",
        "MIMIC_IPW_ESS": f"{float(mimic_diagnostics['effective_sample_size']):,.1f}",
        "EICU_IPW_ESS": f"{float(eicu_diagnostics['effective_sample_size']):,.1f}",
        "EICU_IPW_ESS_ROUNDED": f"{round(float(eicu_diagnostics['effective_sample_size'])):,}",
        "EICU_IPW_AUC": f"{float(eicu_ipw['auroc']):.3f}",
        "CANDIDATE_HOSPITALS": hospital_summary["candidate_hospitals"],
        "STRICT_HOSPITALS": hospital_summary["strict_cohort_hospitals"],
        "BOTH_CLASS_HOSPITALS": hospital_summary["hospitals_with_both_outcome_classes"],
        "FOREST_HOSPITALS": hospital_summary["forest_plot_hospitals"],
        "FOREST_PATIENT_SHARE": f"{100 * float(hospital_summary['forest_plot_patient_share']):.2f}",
        "FOREST_EVENT_SHARE": f"{100 * float(hospital_summary['forest_plot_event_share']):.2f}",
        "CALIBRATION_HOSPITALS": hospital_summary["calibration_hospitals"],
        "NO_ANTIPSYCHOTIC_AUC_CI": metric_ci(no_antipsychotic, "auroc"),
        "NO_ANTIPSYCHOTIC_AUPRC_CI": metric_ci(no_antipsychotic, "auprc"),
        "PSYCHIATRIC_MIMIC_RATE": f"{100 * float(profile['psychiatric_disorder']['mimic_mean']):.1f}",
        "PSYCHIATRIC_EICU_RATE": f"{100 * float(profile['psychiatric_disorder']['eicu_mean']):.1f}",
        "PSYCHIATRIC_SMD": f"{float(profile['psychiatric_disorder']['standardized_mean_difference']):.3f}",
        "MIMIC_MIXED_ICU": get_baseline_value(
            baseline_rows, "ICU type: Mixed or other", "MIMIC-IV"
        ),
        "EICU_MIXED_ICU": get_baseline_value(
            baseline_rows, "ICU type: Mixed or other", "eICU-CRD"
        ),
        "NO_PSYCHIATRIC_AUC_CI": metric_ci(no_psychiatric_external, "auroc"),
        "NO_PSYCHIATRIC_BRIER": f"{float(no_psychiatric_external['brier']):.3f}",
        "NO_PSYCHIATRIC_CAL_INTERCEPT": f"{float(no_psychiatric_external['calibration_intercept']):.3f}",
        "NO_PSYCHIATRIC_CAL_SLOPE": f"{float(no_psychiatric_external['calibration_slope']):.3f}",
        "NO_PSYCHIATRIC_DELTA_FULL": (
            f"{float(coding_sensitivity_differences['psychiatric_disorder']['auroc_difference']):.3f} "
            f"(95% CI "
            f"{float(coding_sensitivity_differences['psychiatric_disorder']['difference_ci_low']):.3f} to "
            f"{float(coding_sensitivity_differences['psychiatric_disorder']['difference_ci_high']):.3f})"
        ),
        "NO_ICU_TYPE_AUC_CI": metric_ci(no_icu_type_external, "auroc"),
        "NO_ICU_TYPE_BRIER": f"{float(no_icu_type_external['brier']):.3f}",
        "NO_ICU_TYPE_CAL_INTERCEPT": f"{float(no_icu_type_external['calibration_intercept']):.3f}",
        "NO_ICU_TYPE_CAL_SLOPE": f"{float(no_icu_type_external['calibration_slope']):.3f}",
        "NO_ICU_TYPE_DELTA_FULL": (
            f"{float(coding_sensitivity_differences['icu_type']['auroc_difference']):.3f} "
            f"(95% CI "
            f"{float(coding_sensitivity_differences['icu_type']['difference_ci_low']):.3f} to "
            f"{float(coding_sensitivity_differences['icu_type']['difference_ci_high']):.3f})"
        ),
        "NO_RACE_AUC_CI": metric_ci(no_race_external, "auroc"),
        "NO_RACE_DELTA_FULL": (
            f"{float(coding_sensitivity_differences['race']['auroc_difference']):.3f} "
            f"(95% CI "
            f"{float(coding_sensitivity_differences['race']['difference_ci_low']):.3f} to "
            f"{float(coding_sensitivity_differences['race']['difference_ci_high']):.3f})"
        ),
        "LOGISTIC_CLINICAL_EXT_AUC_CI": metric_ci(
            logistic_clinical_external, "auroc"
        ),
        "LOGISTIC_ENHANCED_EXT_AUC_CI": metric_ci(
            logistic_enhanced_external, "auroc"
        ),
        "STRICT_ENDPOINT_EXT_AUC_CI": metric_ci(
            strict_enhanced_external, "auroc"
        ),
        "ANY_ENDPOINT_EXT_AUC_CI": metric_ci(
            any_enhanced_external, "auroc"
        ),
        "STRICT_MIMIC_EVENTS": f"{int(endpoint_count_map[('strict_two_positive_days', 'MIMIC')]['events']):,}",
        "STRICT_EICU_EVENTS": f"{int(endpoint_count_map[('strict_two_positive_days', 'eICU')]['events']):,}",
        "ANY_MIMIC_EVENTS": f"{int(endpoint_count_map[('any_post24_delirium', 'MIMIC')]['events']):,}",
        "ANY_EICU_EVENTS": f"{int(endpoint_count_map[('any_post24_delirium', 'eICU')]['events']):,}",
        "RECAL_INTERCEPT_BRIER": f"{float(recalibration['intercept_only']['brier']):.3f}",
        "RECAL_FULL_BRIER": f"{float(recalibration['intercept_and_slope']['brier']):.3f}",
        "RECAL_FULL_ECE": f"{float(recalibration['intercept_and_slope']['ece_10']):.3f}",
        "MISSING_INDICATOR_AUC_CI": metric_ci(missing_indicator, "auroc"),
        "FULL_INDICATOR_AUC_CI": metric_ci(full_indicator, "auroc"),
        "BROAD_AUC": f"{float(broad['auroc']):.3f}",
        "BROAD_AUPRC": f"{float(broad['auprc']):.3f}",
        "BROAD_CAL_INTERCEPT": f"{float(broad['calibration_intercept']):.3f}",
        "BEDSIDE_AUC": f"{float(bedside['auroc']):.3f}",
        "BEDSIDE_BRIER": f"{float(bedside['brier']):.3f}",
        "BEDSIDE_SLOPE": f"{float(bedside['calibration_slope']):.3f}",
        "TRAJECTORY_ENTROPY": f"{float(trajectory_two['entropy']):.3f}",
        "TRAJECTORY_HIGH_N": f"{int(float(trajectory_class['class_n'])):,}",
        "TRAJECTORY_HIGH_RATE": f"{trajectory_rate:.2f}",
        "TRAJECTORY_EXT_AUC": f"{float(trajectory_external['macro_ovr_auroc']):.3f}",
        "POSTGRES_VERSION": "18.4",
        "PYTHON_VERSION": "3.13.12",
        "PANDAS_VERSION": "3.0.5",
        "SKLEARN_VERSION": "1.9.0",
        "XGBOOST_VERSION": "3.3.0",
        "SHAP_VERSION": "0.50.0",
        "R_VERSION": "4.6.0",
        "LCMM_VERSION": "2.2.2",
        "ZENODO_VERSION_URL": "https://doi.org/10.5281/zenodo.21613442",
    }
    return context


TOKEN_PATTERN = re.compile(r"\{\{([A-Z0-9_]+)\}\}")


def replace_tokens(text: str, context: dict[str, str]) -> str:
    def replacement(match: re.Match[str]) -> str:
        key = match.group(1)
        if key == "WORD_COUNT":
            return "{{WORD_COUNT}}"
        if key not in context:
            raise KeyError(f"No value for manuscript token {key}")
        return context[key]

    return TOKEN_PATTERN.sub(replacement, text)


def markdown_table(rows: list[dict[str, str]], columns: list[str] | None = None) -> str:
    if not rows:
        return ""
    columns = columns or list(rows[0].keys())

    def clean(value: object) -> str:
        return str(value if value is not None else "").replace("|", "\\|").replace("\n", " ")

    output = [
        "| " + " | ".join(columns) + " |",
        "| " + " | ".join("---" for _ in columns) + " |",
    ]
    for row in rows:
        output.append("| " + " | ".join(clean(row.get(column, "")) for column in columns) + " |")
    return "\n".join(output)


FIGURES = {
    "[[FIGURE1]]": (
        ASSETS / "figure_1_cohort_flow.png",
        "Figure 1. Cohort formation.",
    ),
    "[[FIGURE2]]": (
        MODELS / "nursing_enhanced_harmonized_shap_summary.png",
        "Figure 2. SHAP summary for the enhanced model.",
    ),
    "[[FIGURE3]]": (
        ASSETS / "figure_3_calibration.png",
        "Figure 3. Calibration of the enhanced model.",
    ),
    "[[FIGURE4]]": (
        MODELS / "eicu_assessment_selection_mechanism.png",
        "Figure 4. Assessment-selection mechanism in eICU-CRD.",
    ),
    "[[FIGURE5]]": (
        MODELS / "nursing_enhanced_harmonized_eicu_hospital_auroc_forest.png",
        "Figure 5. Hospital-level external AUROC.",
    ),
    "[[FIGURES1]]": (
        GBTM / "mimic_trajectory_profiles.png",
        "Supplementary Figure S1. Exploratory daily delirium trajectories.",
    ),
    "[[FIGURES2]]": (
        ASSETS / "figure_s2_decision_curve.png",
        "Supplementary Figure S2. Decision-curve analysis.",
    ),
}


TABLES = {
    "[[TABLE1]]": (
        ASSETS / "table_1_baseline_characteristics.csv",
        "Table 1. Baseline and first-24-hour characteristics of the primary cohorts.",
    ),
    "[[TABLE2]]": (
        ASSETS / "table_2_model_performance.csv",
        "Table 2. Performance of the primary prediction models.",
    ),
    "[[TABLE3]]": (
        ASSETS / "table_3_assessment_selection.csv",
        "Table 3. Selection into the strict delirium-assessment cohort.",
    ),
    "[[TABLES1]]": (
        ASSETS / "table_s1_feature_missingness.csv",
        "Supplementary Table S1. Cross-database feature missingness and transportability profile.",
    ),
    "[[TABLES2]]": (
        ASSETS / "table_s2_coding_harmonization_sensitivity.csv",
        "Supplementary Table S2. Coding-harmonization sensitivity analyses.",
    ),
    "[[TABLES3]]": (
        ASSETS / "table_s3_endpoint_and_logistic_sensitivity.csv",
        "Supplementary Table S3. Alternative endpoint and regularized logistic regression analyses.",
    ),
    "[[TABLES4]]": (
        ASSETS / "table_s4_subgroup_performance.csv",
        "Supplementary Table S4. Performance by sex, age group, and harmonized race category.",
    ),
}


def build_markdown(base_text: str) -> str:
    text = base_text.replace("[[PAGE_BREAK]]", "\n---\n")
    for marker, (path, caption) in TABLES.items():
        rows = read_rows(path)
        text = text.replace(marker, f"**{caption}**\n\n{markdown_table(rows)}")
    for marker, (path, caption) in FIGURES.items():
        relative = path.relative_to(HERE).as_posix() if path.is_relative_to(HERE) else path.relative_to(ROOT).as_posix()
        target = relative if path.is_relative_to(HERE) else f"../{relative}"
        text = text.replace(marker, f"![{caption}]({target})")
    return text


def set_cell_margins(cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    values = {
        "top": CELL_MARGIN_TOP_DXA,
        "bottom": CELL_MARGIN_BOTTOM_DXA,
        "start": CELL_MARGIN_START_DXA,
        "end": CELL_MARGIN_END_DXA,
    }
    for side, value in values.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_indent(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_width.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = True


def add_table(document: Document, path: Path, caption: str) -> None:
    rows = read_rows(path)
    columns = list(rows[0].keys())
    add_caption(document, caption)
    table = document.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_indent(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for index, column in enumerate(columns):
        cell = header.cells[index]
        cell.text = column
        shade_cell(cell, HEADER_FILL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    for source in rows:
        row = table.add_row()
        prevent_row_split(row)
        for index, column in enumerate(columns):
            cell = row.cells[index]
            cell.text = str(source.get(column, ""))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.size = Pt(8)
        if source.get("Section"):
            shade_cell(row.cells[0], "E8EEF5")
            for run in row.cells[0].paragraphs[0].runs:
                run.bold = True

    table_widths = {
        "[[TABLE1]]": [1050, 4050, 1950, 1950],
        "[[TABLE2]]": [1450, 1300, 1275, 1275, 1300, 1100, 1100],
        "[[TABLE3]]": [2250, 1050, 1350, 900, 900, 1650],
        "[[TABLES1]]": [1800, 1100, 1100, 1050, 1100, 1100, 1500],
        "[[TABLES2]]": [1800, 1100, 1250, 1250, 700, 1300, 1960],
        "[[TABLES3]]": [1450, 2450, 1350, 1370, 1370, 1370],
        "[[TABLES4]]": [1250, 1250, 1050, 650, 700, 1900, 2560],
    }
    marker = next(key for key, value in TABLES.items() if value[0] == path)
    widths = table_widths.get(marker)
    if widths and len(widths) == len(columns):
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width / 1440)

    note = document.add_paragraph()
    note.style = document.styles["Table Note"]
    if marker == "[[TABLE1]]":
        note.add_run(
            "Values are median [interquartile range] or n (%). Continuous summaries use "
            "available observations; percentages use the full cohort denominator. "
            "GCS indicates Glasgow Coma Scale; RASS, Richmond Agitation-Sedation Scale; "
            "SpO2, peripheral oxygen saturation; WBC, white blood cell count; BUN, blood urea nitrogen."
        )
    elif marker == "[[TABLE2]]":
        note.add_run(
            "OOF indicates out-of-fold; AUROC, area under the receiver operating characteristic curve; "
            "AUPRC, area under the precision-recall curve. Confidence intervals use 500 bootstrap samples; "
            "external samples used two-stage hospital-and-patient resampling."
        )
    elif marker == "[[TABLES2]]":
        note.add_run(
            "Analyses were post hoc; external AUROC differences used 1,000 "
            "two-stage hospital-and-patient bootstrap samples."
        )
    elif marker == "[[TABLE3]]":
        note.add_run(
            "Selection denotes meeting the strict baseline-negative and longitudinal assessment criteria. "
            "IPW effective sample size is reported for the patient-feature propensity model used in weighting."
        )
    elif marker == "[[TABLES3]]":
        note.add_run(
            "All XGBoost sensitivity models repeated nested tuning. External "
            "confidence intervals used two-stage hospital-and-patient resampling."
        )
    elif marker == "[[TABLES4]]":
        note.add_run(
            "Subgroup estimates were suppressed when fewer than 20 events or "
            "20 non-events were available. Estimates are descriptive and were "
            "not adjusted for multiplicity."
        )


def fit_image(path: Path, max_width_inches: float = 6.25, max_height_inches: float = 7.25) -> tuple[float, float]:
    with Image.open(path) as image:
        width_px, height_px = image.size
    ratio = width_px / height_px
    width = max_width_inches
    height = width / ratio
    if height > max_height_inches:
        height = max_height_inches
        width = height * ratio
    return width, height


def add_figure(document: Document, path: Path, caption: str) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    width, height = fit_image(path)
    paragraph.add_run().add_picture(str(path), width=Inches(width), height=Inches(height))
    if not caption:
        return
    cap = document.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.keep_with_next = False
    cap.add_run(caption).bold = True


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_run_font(run, size: float | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")


def add_inline_runs(paragraph, text: str, size: float | None = None) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, italic=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    heading_specs = {
        "Title": (16, 0, 10),
        "Heading 1": (16, 18, 10),
        "Heading 2": (13, 12, 6),
        "Heading 3": (12, 8, 4),
    }
    for style_name, (size, before, after) in heading_specs.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.05

    caption = document.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(DARK)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True

    table_note = document.styles.add_style("Table Note", 1)
    table_note.font.name = "Calibri"
    table_note._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    table_note.font.size = Pt(8)
    table_note.font.italic = True
    table_note.paragraph_format.space_before = Pt(3)
    table_note.paragraph_format.space_after = Pt(8)
    table_note.paragraph_format.line_spacing = 1.0

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("Delirium prediction transportability")
    set_run_font(run, size=8)
    run.font.color.rgb = RGBColor(110, 118, 126)
    add_page_number(section.footer.paragraphs[0])


def add_reference_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.0
    add_inline_runs(paragraph, text, size=9.5)


def build_docx(base_text: str) -> None:
    document = Document()
    configure_document(document)

    in_references = False
    for raw_line in base_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line == "[[PAGE_BREAK]]":
            document.add_page_break()
            continue
        if line in TABLES:
            path, caption = TABLES[line]
            add_table(document, path, caption)
            continue
        if line in FIGURES:
            path, caption = FIGURES[line]
            if line in {"[[FIGURES1]]", "[[FIGURES2]]"}:
                caption = ""
            add_figure(document, path, caption)
            continue
        if line.startswith("# "):
            paragraph = document.add_paragraph(style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(paragraph, line[2:])
            continue
        if line.startswith("## "):
            heading_text = line[3:]
            document.add_paragraph(heading_text, style="Heading 1")
            in_references = heading_text == "References"
            continue
        if line.startswith("### "):
            document.add_paragraph(line[4:], style="Heading 2")
            continue
        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, line[2:])
            continue
        if in_references and re.match(r"^\d+\.\s", line):
            add_reference_paragraph(document, line)
            continue
        paragraph = document.add_paragraph()
        if line.startswith("**") and ":**" in line:
            paragraph.paragraph_format.keep_with_next = False
        add_inline_runs(paragraph, line)

    document.core_properties.title = (
        "Explainable Prediction of Late or Persistent Delirium in Serially "
        "Assessed Older ICU Patients"
    )
    document.core_properties.subject = "MIMIC-IV development and eICU-CRD transportability assessment"
    document.core_properties.author = "Ren Jingzhuo; Zhang Qiannan; Liu Lixin; Xue Zhaoping"
    document.core_properties.keywords = (
        "delirium; intensive care; machine learning; external validation; transportability"
    )
    document.save(DOCX_OUT)


def main() -> None:
    context = build_context()
    template = TEMPLATE.read_text(encoding="utf-8")
    text = replace_tokens(template, context)
    main_match = re.search(r"## Introduction\n(.*?)\n## Declarations", text, flags=re.S)
    if not main_match:
        raise RuntimeError("Could not identify main-text word-count region")
    word_count = len(re.findall(r"\b[A-Za-z0-9][A-Za-z0-9'-]*\b", main_match.group(1)))
    text = text.replace("{{WORD_COUNT}}", f"{word_count:,}")
    unresolved = TOKEN_PATTERN.findall(text)
    if unresolved:
        raise RuntimeError(f"Unresolved manuscript tokens: {unresolved}")

    markdown = build_markdown(text)
    MARKDOWN_OUT.write_text(markdown, encoding="utf-8")
    build_docx(text)
    print(f"Created {MARKDOWN_OUT}")
    print(f"Created {DOCX_OUT}")
    print(f"Main-text word count: {word_count:,}")


if __name__ == "__main__":
    main()
