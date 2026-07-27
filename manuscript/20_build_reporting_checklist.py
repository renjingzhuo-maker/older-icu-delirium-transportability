from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
SOURCE = HERE / "TRIPOD_AI_CHECKLIST.md"
OUTPUT = HERE / "TRIPOD_AI_CHECKLIST.docx"

PAGE_MAP = {
    "1": "1",
    "2": "2",
    "3a": "3",
    "3b": "3",
    "3c": "3, 20-21, 30-31",
    "4": "3",
    "5a": "4",
    "5b": "4",
    "6a": "4",
    "6b": "4",
    "6c": "5",
    "7": "5, 26",
    "8a": "4",
    "8b": "21",
    "8c": "4",
    "9a": "5",
    "9b": "4-5, 26-28",
    "9c": "21",
    "10": "5-6",
    "11": "5, 26-28",
    "12a": "5-6",
    "12b": "5",
    "12c": "5-6",
    "12d": "6-7",
    "12e": "6",
    "12f": "6, 12",
    "12g": "6, 22",
    "13": "5-6",
    "14": "7, 30-31",
    "15": "6",
    "16": "9-11, 26-28",
    "17": "8, 22",
    "18a": "22",
    "18b": "22",
    "18c": "8, 22",
    "18d": "22",
    "18e": "22",
    "18f": "8, 22",
    "19": "23",
    "20a": "8-9",
    "20b": "9-11, 26-28",
    "20c": "9-11, 26-28",
    "21": "8-17, 26-31",
    "22": "22",
    "23a": "11-12, 29-31",
    "23b": "15-16",
    "24": "12",
    "25": "17-21",
    "26": "20-21",
    "27a": "20",
    "27b": "20-21",
    "27c": "20-21",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def parse_source() -> tuple[str, list[tuple[str, str, str]]]:
    text = SOURCE.read_text(encoding="utf-8")
    study_match = re.search(r"^Study:\s+\*(.+)\*$", text, flags=re.MULTILINE)
    if not study_match:
        raise ValueError("Study title was not found in the checklist source.")
    rows = []
    for line in text.splitlines():
        if not line.startswith("|") or line.startswith("|---") or "| Item |" in line:
            continue
        parts = [part.strip() for part in line.strip("|").split("|")]
        if len(parts) == 3:
            rows.append((parts[0], parts[1], parts[2]))
    return study_match.group(1), rows


def main() -> None:
    study, rows = parse_source()
    missing_pages = sorted(item for item, _, _ in rows if item not in PAGE_MAP)
    if missing_pages:
        raise ValueError(f"Missing page mapping for checklist items: {missing_pages}")

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(8.5)
    normal.paragraph_format.space_after = Pt(3)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TRIPOD+AI Reporting Checklist")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(17)
    run.font.color.rgb = RGBColor(46, 116, 181)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(study)
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(9.5)

    note = doc.add_paragraph(
        "Checklist based on the TRIPOD+AI expanded checklist dated 7 February 2024. "
        "Page references correspond to the 32-page rendered submission draft."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.runs[0].font.size = Pt(8)
    doc.add_page_break()

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    widths = [700, 5900, 5700, 1200]
    headers = ["Item", "Reporting requirement", "Location", "Page(s)"]
    for cell, label, width in zip(table.rows[0].cells, headers, widths):
        set_cell_width(cell, width)
        set_cell_shading(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(label)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(8.5)

    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    prevent_row_split(table.rows[0])

    for item, requirement, location in rows:
        row = table.add_row()
        prevent_row_split(row)
        cells = row.cells
        values = [item, requirement, location, PAGE_MAP[item]]
        for index, (cell, value, width) in enumerate(zip(cells, values, widths)):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if index in {0, 3} else WD_ALIGN_PARAGRAPH.LEFT
            )
            run = paragraph.add_run(value)
            run.font.name = "Arial"
            run.font.size = Pt(8)

    reference = doc.add_paragraph()
    reference.paragraph_format.space_before = Pt(5)
    run = reference.add_run(
        "Reference: Collins GS, Moons KGM, Dhiman P, et al. "
        "TRIPOD+AI statement. BMJ. 2024;385:e078378. "
        "doi:10.1136/bmj-2023-078378."
    )
    run.font.name = "Arial"
    run.font.size = Pt(8)

    properties = doc.core_properties
    properties.title = "TRIPOD+AI Reporting Checklist"
    properties.subject = study
    properties.author = "Jingzhuo Ren"
    doc.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    main()
